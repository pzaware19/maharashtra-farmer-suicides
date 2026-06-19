"""
D1_make_docx.py
Author: Piyush Zaware
Last updated: 2026-06-19

Convert article_draft.md to a formatted Word document for submission to The Print.
Embeds the four figures inline.

OUT
  output/reports/article_farmer_suicides.docx
"""

import os
import re
from pathlib import Path

root   = Path("/Users/piyushzaware/Documents/Unsupervised ML/Maharashtra_Farmer_Suicides")
OUTDIR = root / "output"
FIGDIR = OUTDIR / "figures"
REPDIR = OUTDIR / "reports"

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    import subprocess, sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

# ── Read markdown ──────────────────────────────────────────────────────────────
with open(REPDIR / "article_draft.md", encoding="utf-8") as f:
    raw = f.read()

# ── Build Word doc ─────────────────────────────────────────────────────────────
doc = Document()

# Page margins: 1.3 inches all around
for section in doc.sections:
    section.left_margin   = Inches(1.3)
    section.right_margin  = Inches(1.3)
    section.top_margin    = Inches(1.2)
    section.bottom_margin = Inches(1.2)

# Default body font
style = doc.styles["Normal"]
style.font.name = "Georgia"
style.font.size = Pt(11)

def set_font(run, bold=False, italic=False, size=11, colour=None):
    run.bold   = bold
    run.italic = italic
    run.font.name = "Georgia"
    run.font.size = Pt(size)
    if colour:
        run.font.color.rgb = RGBColor(*colour)

def add_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("_" * 72)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(7)

# Figure insertion order: map, timeseries, vidarbha, causes
fig_sequence = [
    ("fig_map_district.png",     "Figure 1. Average annual farmer suicides by Maharashtra district, 2001-2022. Source: NCRB ADSI."),
    ("fig_timeseries.png",       "Figure 2. Total farmer suicides in Maharashtra, 2001-2022, with key policy events. Source: NCRB ADSI."),
    ("fig_vidarbha_vs_rest.png", "Figure 3. Farmer suicides by region: Vidarbha, Marathwada, and the rest of Maharashtra. Source: NCRB ADSI."),
    ("fig_causes.png",           "Figure 4. Causes of farmer suicides in Maharashtra, 2015, 2019, and 2022. Source: NCRB ADSI Table A-2.4."),
]
fig_index = 0  # insert figures at section breaks

lines = raw.split("\n")
i = 0
while i < len(lines):
    line = lines[i].rstrip()

    # Title
    if line.startswith("# "):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(8)
        run = p.add_run(line[2:])
        set_font(run, bold=True, size=20)
        i += 1
        continue

    # H2 section heading
    if line.startswith("## "):
        # Insert next figure before each section heading (after first)
        if fig_index < len(fig_sequence):
            fname, caption = fig_sequence[fig_index]
            fpath = FIGDIR / fname
            if fpath.exists():
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(10)
                run = p.runs[0] if p.runs else p.add_run()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_picture(str(fpath), width=Inches(5.5))
                cap = doc.add_paragraph(caption)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.paragraph_format.space_after = Pt(14)
                for run in cap.runs:
                    set_font(run, italic=True, size=9, colour=(0x77,0x77,0x77))
            fig_index += 1

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after  = Pt(6)
        run = p.add_run(line[3:])
        set_font(run, bold=True, size=13)
        i += 1
        continue

    # Byline bold
    if line.startswith("**By"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        text = line.strip("*").replace("**", "")
        run = p.add_run(text)
        set_font(run, bold=True, size=10)
        i += 1
        continue

    # Byline italic affiliation
    if line.startswith("*Researcher"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(16)
        text = line.strip("*")
        run = p.add_run(text)
        set_font(run, italic=True, size=10, colour=(0x55,0x55,0x55))
        i += 1
        continue

    # Horizontal rule
    if line.startswith("---"):
        add_rule(doc)
        i += 1
        continue

    # Italic footer note
    if line.startswith("*Piyush Zaware"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        text = line.strip("*")
        run = p.add_run(text)
        set_font(run, italic=True, size=9, colour=(0x77,0x77,0x77))
        i += 1
        continue

    # Empty line
    if line == "":
        i += 1
        continue

    # Regular body paragraph
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(line)
    set_font(run, size=11)
    i += 1

# Append remaining figures at end
while fig_index < len(fig_sequence):
    fname, caption = fig_sequence[fig_index]
    fpath = FIGDIR / fname
    if fpath.exists():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(str(fpath), width=Inches(5.5))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(14)
        for run in cap.runs:
            set_font(run, italic=True, size=9, colour=(0x77,0x77,0x77))
    fig_index += 1

out_path = REPDIR / "article_farmer_suicides.docx"
doc.save(str(out_path))
print(f"Saved: {out_path}  ({out_path.stat().st_size // 1024}K)")
